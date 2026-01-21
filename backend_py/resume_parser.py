"""Resume parsing utilities for PDF and DOCX files."""

import io
from typing import Optional, Tuple
import pdfplumber
from docx import Document


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF file.
    
    Args:
        file_content: Binary content of PDF file
        
    Returns:
        Extracted text from PDF
        
    Raises:
        ValueError: If PDF is corrupted or empty
    """
    try:
        text = ""
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            if len(pdf.pages) == 0:
                raise ValueError("PDF file is empty")
            
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if not text.strip():
            raise ValueError("No text could be extracted from PDF")
        
        return text
    except Exception as e:
        raise ValueError(f"Error parsing PDF: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        file_content: Binary content of DOCX file
        
    Returns:
        Extracted text from DOCX
        
    Raises:
        ValueError: If DOCX is corrupted or empty
    """
    try:
        doc = Document(io.BytesIO(file_content))
        
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Also extract from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        
        if not text.strip():
            raise ValueError("No text could be extracted from DOCX")
        
        return text
    except Exception as e:
        raise ValueError(f"Error parsing DOCX: {str(e)}")


def extract_text_from_resume(file_content: bytes, filename: str) -> Tuple[str, str]:
    """
    Extract text from resume file (PDF or DOCX).
    
    Args:
        file_content: Binary content of resume file
        filename: Name of the file (to determine format)
        
    Returns:
        Tuple of (extracted_text, file_format)
        
    Raises:
        ValueError: If file format is unsupported or extraction fails
    """
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        text = extract_text_from_pdf(file_content)
        return text, "pdf"
    elif filename_lower.endswith(('.docx', '.doc')):
        text = extract_text_from_docx(file_content)
        return text, "docx"
    else:
        raise ValueError(f"Unsupported file format: {filename}")

def clean_resume_text(text: str) -> str:
    """
    Clean and normalize resume text.
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text
    """
    # Remove excessive whitespace
    lines = text.split('\n')
    lines = [line.strip() for line in lines if line.strip()]
    text = '\n'.join(lines)
    
    # Remove multiple spaces
    text = ' '.join(text.split())
    
    return text